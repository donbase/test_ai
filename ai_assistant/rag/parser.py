from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Iterable, List

import os, time, grpc

import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDoc
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from langchain.chains import RetrievalQA
from langchain.prompts import (ChatPromptTemplate,
                               HumanMessagePromptTemplate,
                               SystemMessagePromptTemplate)
import re
from langchain.schema import Document, BaseRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.chat_models.yandex import ChatYandexGPT
from langchain_community.embeddings.yandex import YandexGPTEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.rate_limiters import InMemoryRateLimiter
from readability import Document as ReadabilityDoc
from re import match
import pandas as pd

#from filters import get_blacklist_prompt

# -------------------------------------------------------------------- #
# Globals
# -------------------------------------------------------------------- #
load_dotenv()
_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115 Safari/537.36"
    )
}
_RETRIEVER = None
_rl = InMemoryRateLimiter(requests_per_second=9,
                          check_every_n_seconds=0.05)
_llm = ChatYandexGPT(
    model_uri=f"gpt://{os.getenv('YC_FOLDER_ID')}/yandexgpt-lite/latest",
    api_key=os.getenv("YC_API_KEY"),
    folder_id=os.getenv("YC_FOLDER_ID"),
    temperature=0.0,
    rate_limiter=_rl
)


# -------------------------------------------------------------------- #
# Helpers: scraping & parsing                                           #
# -------------------------------------------------------------------- #
def _scrape_tables_from_url(url: str) -> List[Document]:
    docs: list[Document] = []
    resp = requests.get(url, headers=_HEADERS, timeout=10)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.content, "html.parser")

    for ti, table in enumerate(soup.find_all("table"), start=1):
        cap = table.find("caption")
        title = cap.get_text(strip=True) if cap else None
        if not title:
            for prev in table.find_previous_siblings():
                if match(r"h[2-4]", prev.name):
                    title = prev.get_text(strip=True)
                    break
        title = title or f"Таблица #{ti}"

        df = pd.read_html(str(table), header=0)[0]
        cols = list(df.columns)

        for ri, row in df.iterrows():
            content = f"Название таблицы: {title}; " + "; ".join(
                f"{c}: {row[c]}" for c in cols
            )
            metadata = {
                "source": url,
                "table_index": ti,
                "row_index": ri,
                "column_names": cols,
                "table_title": title,
            }
            docs.append(Document(page_content=content, metadata=metadata))
    return docs


def _parse_docx_tables(path_or_file) -> List[Document]:
    docs: list[Document] = []
    doc = DocxDoc(path_or_file)
    for ti, table in enumerate(doc.tables, start=1):
        tbl_elm = table._element
        parent = tbl_elm.getparent()
        idx = list(parent).index(tbl_elm)
        title = None
        for prev in reversed(parent[:idx]):
            if prev.tag.endswith("p"):
                p = Paragraph(prev, doc)
                if p.text.strip():
                    title = p.text.strip()
                    break
        display_title = title or f"Таблица #{ti}"

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for ri, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            core = "; ".join(f"{h}: {v}" for h, v in zip(headers, cells))
            docs.append(
                Document(
                    page_content=f"Название таблицы: {display_title}\n{core}",
                    metadata={
                        "source": str(path_or_file),
                        "table_index": ti,
                        "row_index": ri,
                        "table_title": display_title,
                        "column_names": headers,
                    },
                )
            )
    return docs


def _parse_pdf_tables(path_or_file) -> List[Document]:
    docs: list[Document] = []
    with pdfplumber.open(path_or_file) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            for ti, table in enumerate(page.extract_tables(), start=1):
                headers = table[0]
                display_title = f"Таблица #{ti} на странице {page_num}"
                for ri, row in enumerate(table[1:], start=1):
                    if all(cell is None or cell.strip() == "" for cell in row):
                        continue
                    row = [cell.strip() if cell else "" for cell in row]
                    core = "; ".join(f"{h}: {v}" for h, v in zip(headers, row))
                    docs.append(
                        Document(
                            page_content=f"Название таблицы: {display_title}\n{core}",
                            metadata={
                                "source": str(path_or_file),
                                "page": page_num,
                                "table_index": ti,
                                "row_index": ri,
                                "table_title": display_title,
                                "column_names": headers,
                            },
                        )
                    )
    return docs


# -------------------------------------------------------------------- #
# Public API                                                           #
# -------------------------------------------------------------------- #
def parse_docs(
        urls: Iterable[str] = (),
        pdf_files: Iterable = (),
        docx_files: Iterable = (),
) -> List[Document]:
    """Собирает документы из URL, PDF, DOCX в список `langchain.schema.Document`."""
    docs: list[Document] = []

    # ------ URL‑ы ----------------------------------------------------
    for url in urls:
        if not url.strip():
            continue
        try:
            resp = requests.get(url, headers=_HEADERS, timeout=10)
            resp.raise_for_status()
            resp.encoding = "utf-8"
            rd = ReadabilityDoc(resp.text)
            clean_html = rd.summary()
            title = rd.short_title()
            soup = BeautifulSoup(clean_html, "html.parser")
            paras = [p.get_text().strip() for p in soup.find_all("p")]
            paras = [p for p in paras if 50 < len(p) < 3000]
            text = "\n".join(paras)
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": url, "title": title},
                )
            )
            docs.extend(_scrape_tables_from_url(url))
        except Exception as exc:
            print(f"⚠️  Не удалось обработать {url}: {exc}")

    # ------ PDF ------------------------------------------------------
    for f in pdf_files:
        name = getattr(f, "name", str(f))
        try:
            stream = f if hasattr(f, "read") else open(f, "rb")
            text_parts: list[str] = []
            with pdfplumber.open(stream) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    txt = page.extract_text()
                    if txt:
                        text_parts.append(txt)
                        docs.append(
                            Document(
                                page_content=f"[стр. {page_num}]\n{txt}",
                                metadata={"source": name, "page": page_num},
                            )
                        )
            stream.seek(0)
            docs.extend(_parse_pdf_tables(stream))
            if not hasattr(f, "read"):
                stream.close()
        except Exception as exc:
            print(f"⚠️  Ошибка при парсинге {name}: {exc}")

    # ------ DOCX -----------------------------------------------------
    for f in docx_files:
        name = getattr(f, "name", str(f))
        try:
            stream = f if hasattr(f, "read") else open(f, "rb")
            tmp = stream
            if hasattr(f, "read"):  # UploadedFile -> bytes -> NamedTemporaryFile
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                tmp.write(f.read())
                tmp.seek(0)

            doc = DocxDoc(tmp.name)
            for i, p in enumerate(doc.paragraphs, start=1):
                if p.text.strip():
                    docs.append(
                        Document(
                            page_content=f"[стр. {i}]\n{p.text}",
                            metadata={"source": name, "page": i},
                        )
                    )
            docs.extend(_parse_docx_tables(tmp.name))
            if hasattr(f, "read"):
                tmp.close()
                Path(tmp.name).unlink(missing_ok=True)
            else:
                stream.close()
        except Exception as exc:
            print(f"⚠️  Ошибка при парсинге {name}: {exc}")

    return docs


# ---------- вспомогалка: дробим список ----------
def build_index(docs, k: int = 10):
    global _RETRIEVER

    splitter = RecursiveCharacterTextSplitter(separators=["\n\n", "\n", ".", "?"], chunk_size=700, chunk_overlap=100)
    chunks = []
    for d in docs:
        for ch in splitter.split_documents([d]):
            page = d.metadata.get("page")
            if page is not None:
                ch.metadata["page"] = page

                tag = f"[стр. {page}]"
                if not ch.page_content.lstrip().startswith(tag):
                    ch.page_content = f"{tag} {ch.page_content}"
            chunks.append(ch)
    texts = [c.page_content for c in chunks]
    metas = [c.metadata for c in chunks]

    emb = YandexGPTEmbeddings(
        api_key=os.getenv("YC_API_KEY"),
        folder_id=os.getenv("YC_FOLDER_ID"),
    )

    vecs = []
    for txt in texts:
        vecs.append(emb.embed_query(txt))
        time.sleep(0.12)

    text_embeddings = list(zip(texts, vecs))

    store = FAISS.from_embeddings(
        text_embeddings,
        embedding=emb,
        metadatas=metas,
        normalize_L2=True,
    )
    _RETRIEVER = store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": k},
    )
    return _RETRIEVER


# --------------------------------------------------------

'''
def llm_blacklist_check(text: str) -> bool:
    system_pt = SystemMessagePromptTemplate.from_template(
        "Ты фильтр контента. У тебя есть список запрещённых тем.\n"
        "Если присланный текст *хоть как‑то* относится к одной из тем, "
        "ответь ровно 'YES'. Иначе ответь ровно 'NO'."
    )
    human_pt = HumanMessagePromptTemplate.from_template(
        "ЗАПРЕЩЁННЫЕ ТЕМЫ:\n{black}\n\nТЕКСТ ДЛЯ ПРОВЕРКИ:\n{txt}"
    )

    prompt = ChatPromptTemplate.from_messages([system_pt, human_pt])
    messages = prompt.format_messages(black=get_blacklist_prompt(), txt=text)

    guard_llm = ChatYandexGPT(
        model_uri=f"gpt://{os.getenv('YC_FOLDER_ID')}/yandexgpt/latest",
        api_key=os.getenv("YC_API_KEY"),
        folder_id=os.getenv("YC_FOLDER_ID"),
        temperature=0.0,
        rate_limiter=_rl
    )

    reply = guard_llm.invoke(messages).content.strip().upper()
    return reply.startswith("Y")
'''

'''
def rewrite_question(question: str,
                     n: int = 1,
                     max_len: int = 120) -> list[str]:
    """
    Перефразирует исходный вопрос так, чтобы он лучше подходил
    для эмбеддингового поиска / BM25.

    Parameters
    ----------
    question : str
        Оригинальный запрос пользователя.
    n : int, default=1
        Сколько вариантов вернуть (1–5).
        При n>1 LLM старается дать разнообразные переформулировки.
    max_len : int, default=120
        Максимальная длина каждой переформулировки (в символах).

    Returns
    -------
    list[str]
        Список из `n` строк. Если фильтр «запрещённых тем» сработал,
        вернётся список с единственным элементом ― исходным вопросом.
    """

    # 1) Проверяем на запрещённые темы
    if llm_blacklist_check(question):
        return [question.strip()]

    # 2) Строим промпт
    sys_msg = SystemMessagePromptTemplate.from_template(
        "Ты ассистент, который преобразует пользовательские вопросы "
        "в короткие поисковые запросы. "
        "Сохраняй смысл, убирай лишние детали, местоимения, эмоции. "
        "Добавляй ключевые термины, если они явно подразумеваются. "
        f"Длина каждого запроса ≤ {max_len} символов."
    )
    human_msg = HumanMessagePromptTemplate.from_template(
        f"Оригинальный вопрос:\n{question}\n\n"
        f"Сформулируй {n} поисковых запроса (каждый с новой строки)."
    )

    prompt = ChatPromptTemplate.from_messages([sys_msg, human_msg])
    messages = prompt.format_messages(max_len=max_len, question=question, n=n)

    # 3) Запрашиваем LLM
    try:
        reply = _llm.invoke(messages).content.strip()
    except Exception as exc:
        # в случае ошибки возвращаем оригинал, чтобы не блокировать пайплайн
        print(f"⚠️  rewrite_question error: {exc}")
        return [question.strip()]

    # 4) Разбираем ответ в список строк
    rewrites = [
        line.lstrip("•-–0123456789. ").rstrip()
        for line in reply.splitlines()
        if line.strip()
    ]

    # 5) Гарантируем, что вернули ровно n вариантов
    if not rewrites:
        return [question.strip()]
    if len(rewrites) < n:
        rewrites += [rewrites[0]] * (n - len(rewrites))
    return rewrites[:n]


# --------------------------------------------------------

# -------------------------------------------------------------------- #
# 1.  mini-retriever «из списка» (нужен для RetrievalQA)
# -------------------------------------------------------------------- #
class StaticRetriever(BaseRetriever):
    docs: List[Document]

    @property
    def _docs(self) -> List[Document]:
        return self.docs

    def _get_relevant_documents(self, query: str, **_) -> List[Document]:
        return self._docs

    async def _aget_relevant_documents(self, query: str, **_):
        return self._docs


# -------------------------------------------------------------------- #
# 2.  вспомогалка для косинусного ранжирования
# -------------------------------------------------------------------- #
_emb = YandexGPTEmbeddings(
    api_key=os.getenv("YC_API_KEY"),
    folder_id=os.getenv("YC_FOLDER_ID")
)


def _cos(a: List[float], b: List[float]) -> float:
    return sum(x * y for x, y in zip(a, b))


def _rerank(question: str, docs: List[Document], top_n: int = 5) -> List[Document]:
    q_vec = _emb.embed_query(question)
    scored = sorted(
        docs,
        key=lambda d: _cos(q_vec, _emb.embed_query(d.page_content)),
        reverse=True,
    )
    return scored[:top_n]


# -------------------------------------------------------------------- #
# 3.  основной entry-point
# -------------------------------------------------------------------- #
def get_answer(question: str) -> dict:
    reject_text = (
        "Извини, но я не могу давать советы на эту тему. Моя задача — помогать "
        "с режимом дня, трекингом сна, кормления, миссий и другими функциями "
        "приложения. Если вопрос важный для здоровья или семьи, лучше "
        "обратиться к специалисту."
    )

    SYSTEM_PT = SystemMessagePromptTemplate.from_template(
        f"""
        "Ты ИИ-чат, созданный для поддержки молодых родителей. Твоя главная задача — предоставлять информацию и помощь в максимально бережном и поддерживающем тоне. Помни, что родители могут испытывать тревогу и неуверенность, поэтому важно создавать атмосферу доверия и безопасности.

        Избегай следующих формулировок и категорий:

        Тревога и паника:
        «Это очень опасно» лучше используй «Важно обратить внимание на…»
        «Срочно бегите к врачу»  лучше используй «Рекомендуется проконсультироваться со специалистом»
        «Тревожный симптом» или «Тревожный признак» лучше используй «Момент, требующий наблюдения»

        Чувство вины:
        «Вы неправильно делаете»  лучше используй «Есть разные подходы к решению»
        «Любой нормальный родитель знает»  лучше используй «Многие родители сталкиваются с этим вопросом»
        «Вы должны»  лучше используй «Возможно, стоит рассмотреть»

        Критика и давление:
        «Вам нужно немедленно изменить»  лучше используй «Можно подумать о вариантах»
        «Все нормальные дети в этом возрасте уже»  лучше используй «Каждый ребёнок развивается в своём темпе»
        «Вы слишком балуете»  лучше используй «Давайте обсудим различные стратегии»

        Обесценивание чувств:
        «Не стоит так переживать»  лучше используй «Ваши переживания абсолютно нормальны»
        «Это ерунда, у других ещё хуже»  лучше используй «Каждый опыт уникален»
        «Вы слишком чувствительны»  лучше используй «Ваши эмоции важны»

        Категоричность:
        «Никогда не делайте так»  лучше используй «Стоит учитывать определённые моменты»
        «Всегда нужно»  лучше используй «Многие родители находят полезным»
        «Единственный правильный способ»  лучше используй «Один из эффективных подходов»

        Формулировки про нормы развития:
        «Ваш ребёнок отстаёт от нормы» лучше используй «Развитие каждого ребёнка имеет свой индивидуальный график»
        «В этом возрасте все дети уже умеют» лучше используй «Многие дети осваивают это умение примерно в таком возрасте»
        «Это не соответствует возрастным нормам» лучше используй «Развитие навыков может происходить в разные сроки»
        «Ваш ребёнок развивается слишком медленно» лучше используй «Темпы развития могут варьироваться»
        «Нужно срочно догонять сверстников» лучше используй «Важно поддерживать естественное развитие ребёнка»
        «Все дети в группе уже умеют, а ваш нет» лучше используй «Каждый ребёнок осваивает навыки в своём темпе»
        «Это признак задержки развития» лучше используй «Стоит обратить внимание на динамику развития»

        Формулировки про паравильно или неправильно:
        «Типичные ошибки родителей» — вместо этого говори о «разных способах решения задач»
        «Правильные подходы» — лучше используй «возможные варианты действий»
        «Неправильный подход» — лучше говори о «разных стратегиях взаимодействия»
        «Нужно обязательно» — лучше используй «можно рассмотреть»
        «Все родители должны» — лучше говори о «разных практиках, которые используют родители»

        При ответах используй:
        Поддержку и понимание
        Конструктивные предложения
        Альтернативные варианты
        Уважение к индивидуальности каждой семьи»"

        Другие правила работы:
        1. Отвечай ТОЛЬКО фактами из CONTEXT. Не добавляй личные знания.
        2. Каждый факт, взятый из документа, снабжай ссылкой [стр. X], сохраняя квадратные скобки,
        Но только в том случае, если данные взяты НЕ с сайта.
        3. Сначала придумай короткий «ПЛАН» размышлений (Thoughts: …), потом
        выведи конечный ответ в секции **Answer**. Пользователю показывается
        только Answer.
        4. Если в проиндексированных материалах нет ответа - отвечай из своих знаний.
        5. Отвечай ТОЛЬКО на то, что спросили. Не добавляй лишних деталей.
        6. При ответе используй не более 5 наиболее релевантных источников.
        """.strip()
    )

    # ---------------- ПРОМПТ ДЛЯ ПЕРВОГО ЧАНКА ----------------
    QUESTION_PROMPT = ChatPromptTemplate.from_messages([
        SYSTEM_PT,
        HumanMessagePromptTemplate.from_template(
            """
            ### User question
            {question}

            ### CONTEXT_BEGIN
            {context_str}
            ### CONTEXT_END

            --- 
            Перед тем как формировать Answer:
            • Сформулируй запрос к базе знаний одной строкой.
            • Подумай по шагам, какие факты из CONTEXT нужны.
            Напиши их в секции Thoughts, затем выдай секцию Answer.
            Формат:
            Thoughts: <твои размышления, 2–4 предложения>
            Answer: <конечный ответ для пользователя>
            """
        ),
    ])

    REFINE_PROMPT = ChatPromptTemplate.from_messages([
        SYSTEM_PT,
        HumanMessagePromptTemplate.from_template(
            """
            ### Existing draft
            {existing_answer}
    
            ### NEW_CONTEXT_BEGIN
            {context_str}
            ### NEW_CONTEXT_END
    
            Проверь, есть ли в новом контексте сведения, улучшающие точность.
            • Если да – обнови Answer, добавив факты + ссылки.
            • Если нет – верни existing_answer без изменений.
    
            Обязательно сохрани формат:
            Thoughts: ...
            Answer: ...
            """
        ),
    ])

    # --- 3.1. фильтр запрещённого контента --------------------------------
    if llm_blacklist_check(question):
        return {"result": reject_text, "source_documents": []}

    # --- 3.2. multi-query rewrite ----------------------------------------
    rewrites = rewrite_question(question, n=3)  # <-- новая функция

    # --- 3.3. расширенный пул документов ---------------------------------
    pool: List[Document] = []
    seen = set()
    for rq in rewrites:
        for doc in _RETRIEVER.get_relevant_documents(rq, k=10):
            uid = (doc.metadata.get("source"), doc.metadata.get("page"),
                   doc.page_content[:120])
            if uid not in seen:
                pool.append(doc)
                seen.add(uid)

    # --- 3.4. реранжируем и берём максимум 5 -----------------------------
    top_docs = _rerank(question, pool, top_n=5)

    # --- 3.5. готовим retriever «из списка» ------------------------------
    static_ret = StaticRetriever(docs=top_docs)

    qa = RetrievalQA.from_chain_type(
        llm=_llm,
        chain_type="refine",
        retriever=static_ret,
        return_source_documents=True,
        chain_type_kwargs={
            "question_prompt": QUESTION_PROMPT,
            "refine_prompt": REFINE_PROMPT,
        },
    )

    result = qa.invoke(question)

    # --- 3.6. вычищаем «Thoughts: …» -------------------------------------
    result["result"] = re.sub(
        r"Thoughts:.*?Answer:",
        "",
        result["result"],
        flags=re.S
    ).strip()

    # гарантируем, что в ответе не более 5 источников
    result["source_documents"] = top_docs
    return result
'''